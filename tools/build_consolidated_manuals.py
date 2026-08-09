"""Produce clinician-facing consolidated SEMANTIAR manuals (HTML and DOCX)."""
import html
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
HTML_SOURCE = DOCS / "Manual_base_spans.html"
DOCX_SOURCE = DOCS / "Manual_base_spans.docx"
HTML_TARGET = DOCS / "Manual_del_anotador_SemantIAr_v2_consolidado.html"
DOCX_TARGET = DOCS / "Manual_del_anotador_SemantIAr_v2_consolidado.docx"
SCREENSHOT = DOCS / "captura_revision_formas_breves.png"

NAV = [
    ("como-usar-este-manual", "Cómo usar este manual"),
    ("3-que-hara-como-anotador", "Flujo de trabajo"),
    ("4-como-revisar-una-nota-de-principio-a-fin", "Revisar una nota"),
    ("5-como-completar-una-tarjeta-lexica", "Formas breves"),
    ("8-como-revisar-la-informacion-clinica", "Spans y conceptos"),
    ("11-cierre-exportacion-y-problemas-frecuentes", "Cierre y exportación"),
    ("actualizacion-interfaz-20260728", "Flujo actual del anotador web"),
]

RELATED_LINKS = {
    "como-usar-este-manual": [
        ("1-1-lexico-clinica-y-contexto-tres-preguntas-relacionadas", "Léxico, clínica y contexto"),
        ("4-como-revisar-una-nota-de-principio-a-fin", "Recorrido completo de una nota"),
        ("12-guia-rapida-durante-la-anotacion", "Guía rápida"),
    ],
    "1-que-significa-lexico-en-este-manual": [
        ("2-que-es-una-forma-breve", "Qué cuenta como forma breve"),
        ("8-como-revisar-la-informacion-clinica", "Revisión clínica independiente"),
    ],
    "2-que-es-una-forma-breve": [
        ("5-como-completar-una-tarjeta-lexica", "Completar una tarjeta léxica"),
        ("anexo-a-equivalencias-tecnicas-de-la-interfaz", "Equivalencias de los campos"),
    ],
    "3-que-hara-como-anotador": [
        ("4-como-revisar-una-nota-de-principio-a-fin", "Orden de revisión"),
        ("11-1-antes-de-cerrar", "Control antes del cierre"),
    ],
    "4-como-revisar-una-nota-de-principio-a-fin": [
        ("4-3-que-significan-los-resaltados", "Interpretar los resaltados"),
        ("actualizacion-interfaz-20260728", "Interfaz de revisión visible"),
        ("11-cierre-exportacion-y-problemas-frecuentes", "Cierre y exportación"),
    ],
    "5-como-completar-una-tarjeta-lexica": [
        ("6-como-decidir-el-significado", "Decidir el significado"),
        ("7-funcion-seccion-pistas-y-comentario", "Función, sección y pistas"),
    ],
    "6-como-decidir-el-significado": [
        ("6-1-opciones-de-decision", "Opciones de decisión"),
        ("sentido-resuelto", "Cuándo usar Sentido resuelto"),
        ("11-1-antes-de-cerrar", "Validación antes de cerrar"),
    ],
    "7-funcion-seccion-pistas-y-comentario": [
        ("7-2-seccion-donde-aparece", "Sección local de la nota"),
        ("7-3-pistas-por-que-tomo-la-decision", "Pistas contextuales"),
        ("anexo-b-codigos-de-pistas-contextuales", "Códigos de pistas"),
    ],
    "8-como-revisar-la-informacion-clinica": [
        ("8-1-que-es-una-seleccion-textual-o-span", "Qué es un span"),
        ("8-4-elegir-categoria-y-concepto", "Categoría y concepto"),
        ("spans-superpuestos", "Spans superpuestos"),
        ("spans-discontinuos", "Política de spans discontinuos"),
    ],
    "9-ejemplos-concretos": [
        ("5-como-completar-una-tarjeta-lexica", "Aplicar los ejemplos a la tarjeta"),
        ("8-2-como-elegir-el-limite", "Elegir el límite del span"),
    ],
    "10-caso-de-practica-de-alta-densidad": [
        ("10-1-como-abordarlo", "Cómo abordarlo"),
        ("10-2-ejemplos-de-decisiones-defendibles", "Decisiones defendibles"),
    ],
    "11-cierre-exportacion-y-problemas-frecuentes": [
        ("11-1-antes-de-cerrar", "Lista de control"),
        ("11-3-exportacion", "Exportar el trabajo"),
        ("11-4-si-algo-no-funciona", "Resolver problemas"),
    ],
    "12-guia-rapida-durante-la-anotacion": [
        ("4-2-orden-recomendado", "Orden detallado"),
        ("11-1-antes-de-cerrar", "Verificación final"),
    ],
    "anexo-a-equivalencias-tecnicas-de-la-interfaz": [
        ("5-como-completar-una-tarjeta-lexica", "Uso clínico de los campos"),
        ("actualizacion-interfaz-20260728", "Comportamiento actual de los desplegables"),
    ],
    "anexo-b-codigos-de-pistas-contextuales": [
        ("7-3-pistas-por-que-tomo-la-decision", "Cómo elegir pistas"),
        ("7-4-comentario-solo-cuando-agrega-valor", "Cuándo agregar comentario"),
    ],
}

WORD_BOOKMARKS = {
    "Cómo usar este manual": "xref_manual",
    "1. Qué significa «léxico» en este manual": "xref_lexico",
    "2. Qué es una forma breve": "xref_forma",
    "3. Qué hará como anotador": "xref_flujo",
    "4. Cómo revisar una nota de principio a fin": "xref_revision",
    "5. Cómo completar una tarjeta léxica": "xref_tarjeta",
    "6. Cómo decidir el significado": "xref_sentido",
    "7. Función, sección, pistas y comentario": "xref_contexto",
    "8. Cómo revisar la información clínica": "xref_clinica",
    "9. Ejemplos concretos": "xref_ejemplos",
    "10. Caso de práctica de alta densidad": "xref_practica",
    "11. Cierre, exportación y problemas frecuentes": "xref_cierre",
    "12. Guía rápida durante la anotación": "xref_guia",
    "Anexo A. Equivalencias técnicas de la interfaz": "xref_anexo_a",
    "Anexo B. Códigos de pistas contextuales": "xref_anexo_b",
    "Flujo actual del anotador web: selección, clasificación y cierre": "xref_actualizacion",
}

WORD_RELATED = {
    "Cómo usar este manual": [("Revisión paso a paso", "xref_revision"), ("Guía rápida", "xref_guia")],
    "1. Qué significa «léxico» en este manual": [("Formas breves", "xref_forma"), ("Revisión clínica", "xref_clinica")],
    "2. Qué es una forma breve": [("Tarjeta léxica", "xref_tarjeta"), ("Campos de la interfaz", "xref_anexo_a")],
    "3. Qué hará como anotador": [("Recorrido de una nota", "xref_revision"), ("Cierre", "xref_cierre")],
    "4. Cómo revisar una nota de principio a fin": [("Tarjeta léxica", "xref_tarjeta"), ("Interfaz actualizada", "xref_actualizacion")],
    "5. Cómo completar una tarjeta léxica": [("Decidir el significado", "xref_sentido"), ("Función, sección y pistas", "xref_contexto")],
    "6. Cómo decidir el significado": [("Función, sección y pistas", "xref_contexto"), ("Control de cierre", "xref_cierre")],
    "7. Función, sección, pistas y comentario": [("Códigos de pistas", "xref_anexo_b"), ("Interfaz actualizada", "xref_actualizacion")],
    "8. Cómo revisar la información clínica": [("Ejemplos", "xref_ejemplos"), ("Spans superpuestos", "xref_actualizacion")],
    "9. Ejemplos concretos": [("Tarjeta léxica", "xref_tarjeta"), ("Revisión clínica", "xref_clinica")],
    "10. Caso de práctica de alta densidad": [("Orden de revisión", "xref_revision"), ("Verificación final", "xref_cierre")],
    "11. Cierre, exportación y problemas frecuentes": [("Guía rápida", "xref_guia"), ("Volver al recorrido", "xref_revision")],
    "12. Guía rápida durante la anotación": [("Recorrido detallado", "xref_revision"), ("Control de cierre", "xref_cierre")],
    "Anexo A. Equivalencias técnicas de la interfaz": [("Uso de la tarjeta", "xref_tarjeta"), ("Interfaz actualizada", "xref_actualizacion")],
    "Anexo B. Códigos de pistas contextuales": [("Función, sección y pistas", "xref_contexto"), ("Guía rápida", "xref_guia")],
    "Flujo actual del anotador web: selección, clasificación y cierre": [("Revisar una nota", "xref_revision"), ("Formas breves", "xref_tarjeta"), ("Información clínica", "xref_clinica"), ("Cierre", "xref_cierre")],
}

UPDATE_HTML = """
<section class="manual-section" id="actualizacion-interfaz-20260728" aria-labelledby="actualizacion-titulo">
  <h1 id="actualizacion-titulo">Flujo actual del anotador web: selección, clasificación y cierre</h1>
  <p>Esta edición está dirigida a profesionales de salud que anotan notas clínicas en <strong>SemantIAr Anotador</strong>. La nota queda visible mientras se trabaja: el texto clínico permanece en la columna izquierda y el detalle secuencial se abre en la columna derecha. La interfaz actual presenta una decisión unificada para cada marca.</p>
  <p>Este manual explica cómo ejecutar la tarea. Para el fundamento metodológico —reglas, reconocimiento de entidades, spans, normalización, mapeo clínico, preanotación y criterios de calidad— consulte <a href="Fundamentos_metodologicos_para_anotadores_SemantIAr.docx">Fundamentos metodológicos para anotadores SemantIAr, versión 1.2 (9 de agosto de 2026)</a>.</p>
  <p class="related-links"><strong>Recorrido relacionado:</strong> <a href="#3-que-hara-como-anotador">responsabilidades del anotador</a> · <a href="#4-como-revisar-una-nota-de-principio-a-fin">revisión de una nota</a> · <a href="#8-como-revisar-la-informacion-clinica">spans y conceptos</a> · <a href="#11-cierre-exportacion-y-problemas-frecuentes">cierre y exportación</a></p>
  <h2 id="flujo-tres-pasos">Recorrido en tres pasos</h2>
  <ol>
    <li><strong>Cargar JSON.</strong> Cargue el archivo asignado y elija la nota. Los manuales no contienen notas clínicas ni sirven para editar el JSON a mano.</li>
    <li><strong>Paso 1 · lectura de la nota.</strong> Lea la nota clínica completa antes de marcar. Puede registrar la <em>especialidad sospechada de la nota</em> en el desplegable alfabético; es una pista general y no reemplaza la sección local de cada aparición.</li>
    <li><strong>Paso 2 · marcación de pendientes.</strong> Revise los premarcados, ajuste sus límites y seleccione en el texto toda mención que falte. Después de elegir el tramo exacto, use <strong>“Marcar para revisar”</strong>. Las premarcaciones son candidatos, no respuestas; la lista de decisiones se abre recién al pulsar <strong>“Terminé la marcación · ver pendientes”</strong>.</li>
    <li><strong>Paso 3 · decisión de menciones.</strong> La interfaz muestra una marca por vez. Complete su decisión y avance con <strong>“Siguiente”</strong> o con los controles inferiores de anterior/siguiente; no necesita volver al inicio de la nota.</li>
  </ol>
  <h2 id="clasificacion-unificada">Una clasificación para cada aparición</h2>
  <p>En “¿Qué representa aquí?” elija una sola opción inicial. La misma marca puede conservar dos dimensiones cuando corresponda; no hay que marcar dos veces el mismo texto.</p>
  <table><thead><tr><th>Etiqueta visible</th><th>Qué significa</th><th>Qué se completa</th></tr></thead><tbody>
    <tr><td><strong>Término con información clínica</strong></td><td>La aparición expresa información clínica, aunque todavía no exista un concepto SNOMED CT defendible.</td><td>Jerarquía clínica, búsqueda de concepto, literal y contexto clínico. El estado puede quedar pendiente de codificación.</td></tr>
    <tr><td><strong>Término sin información clínica</strong></td><td>La aparición tiene un significado local o una función léxica, pero no se registra como concepto clínico.</td><td>Tipo de forma, “¿Qué significa aquí?”, decisión, función, sección y pistas.</td></tr>
    <tr><td><strong>Ambos</strong></td><td>La misma aparición tiene información clínica y también requiere una decisión de significado local.</td><td>Complete las dos dimensiones en la misma marca, de forma secuencial.</td></tr>
    <tr><td><strong>No anotar</strong></td><td>La aparición no es anotable según el protocolo.</td><td>No se abren campos de codificación ni de forma breve.</td></tr>
  </tbody></table>
  <p>Si cambia de opción, la interfaz reconfigura la marca existente. <strong>Ambos</strong> conserva o agrega las dos dimensiones; pasar a una sola elimina únicamente la dimensión que ya no corresponde. No cree una segunda marca para compensarlo.</p>
  <h2 id="detalle-secuencial">Detalle secuencial de la marca</h2>
  <p>Al abrir una marca, el panel de detalle queda junto a la nota y termina con <strong>“Mención anterior”</strong>, <strong>“Volver a marcas”</strong> y <strong>“Siguiente mención”</strong>. Use “Volver a marcas” solo para regresar a la lista; la decisión normal continúa hacia abajo con el control siguiente.</p>
  <p>Para una marca clínica, abra la codificación y complete categoría, búsqueda, concepto, texto literal y las cuatro dimensiones de contexto: polaridad, certeza, temporalidad y sujeto. Para una marca sin información clínica, complete la tarjeta léxica. En “¿Qué significa aquí?” el valor elegido o escrito es el que se incorpora a la anotación; no hay un campo adicional que duplique esa decisión.</p>
  <p>Los campos cerrados usan listas desplegables. Al abrir una lista, cada opción muestra su renglón completo y la explicación; “¿Qué pistas usaste?” admite varias selecciones. El único texto libre previsto para la tarjeta es el significado cuando no existe una opción disponible y el comentario breve de trazabilidad si realmente agrega información.</p>
  <h2 id="cierre-flujo-actual">Cierre y exportación</h2>
  <p>Finalice la nota solo cuando no queden marcas pendientes y todas las dimensiones elegidas estén completas. La confirmación de cierre aparece al final de la nota. Descargue el JSON de avance para conservar o entregar el trabajo; para continuar, cargue la copia más reciente desde la interfaz y no la edite manualmente.</p>
  <p class="related-links"><strong>Antes de cerrar:</strong> <a href="#11-1-antes-de-cerrar">lista de control</a> · <a href="#11-3-exportacion">exportación</a> · <a href="#11-4-si-algo-no-funciona">resolver problemas</a></p>
  <h2 id="compatibilidad-movil">Compatibilidad móvil</h2>
  <p>Los campos de escritura rápida para agregar una mención o una forma breve se mantienen como ayuda para celular. En la versión web de escritorio, seleccione el texto directamente en la nota y use “Marcar para revisar”; no duplique la misma aparición desde esos campos.</p>
  <figure class="manual-figure"><img src="captura_revision_formas_breves.png" alt="Captura de referencia de la revisión de formas breves, con la nota clínica fija a la izquierda y los campos a la derecha."><figcaption>Captura usada para revisar el diseño. En la versión corregida, las explicaciones largas se leen al abrir la lista y no quedan concatenadas en el valor seleccionado.</figcaption></figure>
  <h2 id="spans-superpuestos">Spans que se superponen</h2>
  <p>Dos menciones válidas pueden compartir parte del mismo texto. La aplicación ahora las conserva y las muestra con un resaltado rayado y un número. Seleccione cada marca para revisarla por separado. No cambie los límites solo para evitar una superposición; aplique la <a href="#8-como-revisar-la-informacion-clinica">regla de anclaje del span</a> y conserve cada mención clínicamente defendible. No cree superposiciones únicamente para representar niveles de granularidad de una misma mención.</p>
  <h2 id="spans-discontinuos">Política para spans discontinuos</h2>
  <p>La versión actual de SemantIAr utiliza spans continuos como modalidad predeterminada. Los spans discontinuos se registran como casos candidatos, pero no deben construirse desde la interfaz ni forzarse dentro de un único resaltado.</p>
  <p>Si la mención puede representarse con un tramo continuo sin incluir palabras ajenas, use un span continuo. Si la expresión parece discontinua pero puede dividirse en menciones clínicas independientes, sepárelas solo si cada una tiene sentido propio. Si un span continuo incorporaría material que no pertenece a la mención y dividirla distorsionaría el significado, marque el caso para adjudicación o absténgase según el protocolo.</p>
  <p>Esta decisión es gradual: los spans discontinuos pueden ser más fieles al texto, pero requieren una interfaz de selección múltiple, un formato de offsets diferente y reglas adicionales para solapamientos, adjudicación, métricas y modelos. La política se apoya en la revisión sistemática de 44 estudios de Alhassan et al. (2025), en el modelo específico de Dai et al. (2020) y en la evaluación de reconocimiento y normalización de Trivedi et al. (2020). La evidencia y las referencias completas están en <a href="Fundamentos_metodologicos_para_anotadores_SemantIAr.docx">Fundamentos metodológicos para anotadores SemantIAr</a>.</p>
  <aside class="callout key" role="note"><h2>Regla práctica</h2><p>No incluya palabras intermedias que no pertenecen a la entidad solo para evitar una discontinuidad. Registrar el caso para revisión es preferible a alterar el literal o introducir ruido semántico.</p></aside>
  <h2 id="sentido-resuelto">Qué significa “Sentido resuelto”</h2>
  <p>Use “Sentido resuelto” únicamente después de elegir un significado disponible. Si la interfaz no ofrece un significado que pueda confirmar, use “Proponer sentido nuevo”, “Ambigua aun con contexto” o “No puedo determinarla”, según corresponda. Una decisión resuelta sin significado vuelve a <strong>Pendiente</strong> al cargar el archivo y no permite completar la revisión. Consulte <a href="#6-como-decidir-el-significado">Cómo decidir el significado</a> y la lista de <a href="#11-1-antes-de-cerrar">verificación antes del cierre</a>.</p>
  <aside class="callout key" role="note"><h2>Regla práctica</h2><p>Primero confirme el texto y su contexto; después clasifique la marca y complete las dimensiones que correspondan. No fuerce una resolución para cerrar la nota.</p></aside>
  <p><a href="#manual-top">Volver al inicio</a> · <a href="#flujo-tres-pasos">Volver al recorrido en tres pasos</a> · <a href="#clasificacion-unificada">Volver a las etiquetas</a></p>
</section>
"""


def build_html():
    source = HTML_SOURCE.read_text(encoding="utf-8")

    def slugify(value):
        value = re.sub(r"<[^>]+>", "", value)
        value = html.unescape(value).replace("«", "").replace("»", "")
        value = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
        value = re.sub(r"[^a-zA-Z0-9]+", "-", value.lower()).strip("-")
        return value

    def add_heading_id(match):
        open_tag, content, close_tag = match.group(1), match.group(2), match.group(3)
        if re.search(r"\\bid=", open_tag, flags=re.IGNORECASE):
            return match.group(0)
        return f'{open_tag[:-1]} id="{slugify(content)}">{content}{close_tag}'

    source = re.sub(r"(<h[12][^>]*>)(.*?)(</h[12]>)", add_heading_id, source, flags=re.IGNORECASE | re.DOTALL)
    style = """
<style>
html{background:#e9edf2}body{box-sizing:border-box;max-width:920px;margin:0 auto!important;padding:36px 48px 64px!important;background:#fff;color:#202936;font-family:Calibri,Arial,sans-serif!important;font-size:17px!important;line-height:1.6!important;word-wrap:normal!important}body p,body li{font-size:17px!important;line-height:1.6!important}body p span,body li span{font-size:inherit!important;line-height:inherit!important}body h1{margin-top:42px!important;font-size:28px!important;line-height:1.25!important;color:#1f4f7d}body h2{margin-top:28px!important;font-size:21px!important;line-height:1.3!important;color:#2d3e5f}body h1 span,body h2 span{font-size:inherit!important;line-height:inherit!important}.WordSection1{width:auto!important;margin:0!important}.manual-nav{display:none!important}.manual-contents{display:block!important;margin:22px 0 32px!important;padding:20px 22px!important;border:1px solid #9dbde1;border-left:5px solid #2f6fae;border-radius:6px;background:#f4f8fd;font:600 17px/1.55 Calibri,Arial,sans-serif!important}.manual-contents strong{display:block;margin-bottom:10px;color:#1e3a5f;font-size:19px}.manual-contents a{display:block!important;width:max-content;max-width:100%;margin:5px 0!important;color:#075985!important;text-decoration:underline!important}.manual-contents a:hover,.manual-contents a:focus{background:#eaf3fb}.manual-section{scroll-margin-top:1rem}body img{display:block!important;box-sizing:border-box;max-width:100%!important;width:auto!important;height:auto!important;margin:18px auto!important}body table{box-sizing:border-box;max-width:100%!important;width:100%!important;margin:18px 0!important;border-collapse:collapse!important;font-size:15px!important;line-height:1.45!important}body td,body th{font-size:15px!important;line-height:1.45!important;vertical-align:top!important;white-space:normal!important;word-break:normal!important}body td span,body th span{font-size:inherit!important;line-height:inherit!important}@media (max-width:720px){html{background:#fff}body{max-width:none;padding:22px 18px 42px!important;font-size:16px!important}body p,body li{font-size:16px!important}body h1{font-size:25px!important}body h2{font-size:20px!important}body table{font-size:14px!important}body td,body th{font-size:14px!important;padding:7px!important}}
.manual-figure{margin:24px auto;padding:12px;border:1px solid #c8d8eb;background:#f8fbfe}.manual-figure figcaption{margin-top:10px;color:#53627a;font-size:15px;line-height:1.45;text-align:center}
.related-links{margin:8px 0 22px!important;padding:10px 12px!important;border-left:3px solid #9dbde1;background:#f7faff;color:#41516a;font-size:15px!important;line-height:1.5!important}.related-links a{color:#075985!important;text-decoration:underline!important}.related-links strong{color:#2d3e5f}
</style>"""
    nav_links = "".join(f'<a href="#{anchor}">{html.escape(label)}</a>' for anchor, label in NAV)
    nav = f'<nav class="manual-nav" aria-label="Navegación rápida">{nav_links}<a class="manual-nav-top" href="#manual-top">Inicio ↑</a></nav>'
    contents = f'<nav class="manual-contents" aria-label="Contenido navegable"><strong>Ir directamente a una sección</strong>{nav_links}</nav>'
    source = source.replace("</head>", style + "</head>", 1)
    source = re.sub(
        r"<body(\s[^>]*)?>",
        lambda match: match.group(0) + '<a id="manual-top"></a>' + nav,
        source,
        count=1,
        flags=re.IGNORECASE,
    )
    source = re.sub(
        r'(<h1[^>]*id="como-usar-este-manual"[^>]*>)',
        lambda match: contents + match.group(1),
        source,
        count=1,
        flags=re.IGNORECASE,
    )
    for anchor, links in RELATED_LINKS.items():
        related = " · ".join(
            f'<a href="#{target}">{html.escape(label)}</a>' for target, label in links
        )
        paragraph = f'<p class="related-links"><strong>También puede consultar:</strong> {related}</p>'
        source = re.sub(
            rf'(<h1[^>]*id="{re.escape(anchor)}"[^>]*>.*?</h1>)',
            lambda match, paragraph=paragraph: match.group(1) + paragraph,
            source,
            count=1,
            flags=re.IGNORECASE | re.DOTALL,
        )
    source = source.replace("</body>", UPDATE_HTML + "</body>", 1)
    HTML_TARGET.write_text(source, encoding="utf-8")


def build_docx():
    document = Document(DOCX_SOURCE)
    document.add_page_break()
    document.add_heading("Flujo actual del anotador web: selección, clasificación y cierre", level=1)
    document.add_paragraph(
        "Esta edición está dirigida a profesionales de salud que anotan notas clínicas en SemantIAr Anotador. "
        "La nota queda visible mientras se trabaja: el texto clínico permanece en la columna izquierda y el "
        "detalle secuencial se abre en la columna derecha. La interfaz actual presenta una decisión unificada "
        "para cada marca."
    )
    document.add_paragraph(
        "Este documento explica cómo ejecutar la tarea. Para el fundamento metodológico sobre reglas, reconocimiento "
        "de entidades, spans, normalización, mapeo clínico, preanotación y criterios de calidad, consulte "
        "Fundamentos metodológicos para anotadores SemantIAr, versión 1.2 (9 de agosto de 2026)."
    )
    document.add_paragraph(
        "Recorrido en tres pasos: 1) cargue el JSON y elija la nota; 2) lea la nota completa, revise los premarcados "
        "y seleccione las menciones faltantes con “Marcar para revisar”; 3) después de “Terminé la marcación · ver "
        "pendientes”, decida cada marca de forma secuencial. La especialidad sospechada de la nota se elige en un "
        "desplegable alfabético y funciona como pista general; la sección de cada aparición sigue siendo local."
    )
    document.add_paragraph(
        "En “¿Qué representa aquí?” elija una sola opción inicial: “Término con información clínica”, “Término sin "
        "información clínica”, “Ambos” o “No anotar”. “Ambos” conserva las dos dimensiones en la misma marca y no "
        "requiere volver a seleccionarla. Cambiar de opción reconfigura la marca existente; no cree un duplicado."
    )
    document.add_heading("Detalle secuencial de la marca", level=2)
    document.add_paragraph(
        "Al abrir una marca, el panel queda junto a la nota y termina con “Mención anterior”, “Volver a marcas” y "
        "“Siguiente mención”. Para una marca clínica complete jerarquía, búsqueda de concepto, literal y polaridad, "
        "certeza, temporalidad y sujeto. Para una marca sin información clínica complete tipo de forma, “¿Qué "
        "significa aquí?”, decisión, función, sección y pistas. En “¿Qué significa aquí?” el valor elegido o escrito "
        "es el que se incorpora a la anotación; no hay un campo adicional que duplique esa decisión."
    )
    document.add_paragraph(
        "Los campos cerrados usan listas desplegables. Al abrir una lista, cada opción muestra su renglón completo y "
        "la explicación; “¿Qué pistas usaste?” admite varias selecciones. El único texto libre previsto para la tarjeta "
        "es el significado cuando no existe una opción disponible y el comentario breve de trazabilidad cuando agrega "
        "información."
    )
    document.add_heading("Clasificación unificada", level=2)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Etiqueta visible"
    hdr[1].text = "Qué significa"
    hdr[2].text = "Qué se completa"
    rows = [
        ("Término con información clínica", "La aparición expresa información clínica, aunque todavía no exista un concepto SNOMED CT defendible.", "Jerarquía, búsqueda, literal y contexto clínico; puede quedar pendiente de codificación."),
        ("Término sin información clínica", "Tiene significado local o función léxica, pero no se registra como concepto clínico.", "Tipo, significado, decisión, función, sección y pistas."),
        ("Ambos", "La misma aparición tiene información clínica y también requiere significado local.", "Complete las dos dimensiones en la misma marca."),
        ("No anotar", "La aparición no es anotable según el protocolo.", "No se abren campos de codificación ni de forma breve."),
    ]
    for label, meaning, fields in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = label, meaning, fields
    linked = document.add_paragraph()
    linked.add_run("Vínculos de esta sección: ").bold = True
    for index, (label, anchor) in enumerate(
        [("revisión de una nota", "xref_revision"), ("formas breves", "xref_tarjeta"), ("información clínica", "xref_clinica"), ("cierre", "xref_cierre")]
    ):
        if index:
            linked.add_run(" · ")
        add_internal_hyperlink(linked, label, anchor)
    document.add_heading("Cierre y exportación", level=2)
    document.add_paragraph(
        "Finalice la nota solo cuando no queden marcas pendientes y todas las dimensiones estén completas. La "
        "confirmación aparece al final de la nota. Descargue el JSON de avance para conservar "
        "o entregar el trabajo; para continuar, cargue la copia más reciente desde la interfaz y no la edite manualmente."
    )
    screenshot = document.add_picture(str(SCREENSHOT), width=Inches(6.1))
    screenshot._inline.docPr.set(
        "descr",
        "Interfaz de revisión de formas breves con la nota clínica visible a la izquierda y los campos de decisión a la derecha.",
    )
    screenshot._inline.docPr.set("title", "Revisión de formas breves")
    document.add_paragraph(
        "Captura de referencia del diseño de desplegables. En la versión corregida, las explicaciones largas se leen "
        "al abrir la lista y no quedan concatenadas en el valor seleccionado."
    )
    document.add_heading("Spans que se superponen", level=2)
    document.add_paragraph(
        "Dos menciones válidas pueden compartir parte del mismo texto. La aplicación ahora conserva esos spans "
        "y los muestra con un resaltado rayado y una marca numerada. Seleccione cada marca para revisarla por "
        "separado. No modifique los límites solo para evitar la superposición: conserve la mínima mención suficiente "
        "que represente cada mención clínica. No cree superposiciones únicamente para representar niveles de "
        "granularidad de una misma mención."
    )
    document.add_heading("Política para spans discontinuos", level=2)
    document.add_paragraph(
        "La versión actual de SemantIAr utiliza spans continuos como modalidad predeterminada. Los spans discontinuos "
        "se registran como casos candidatos, pero no deben construirse desde la interfaz ni forzarse dentro de un único resaltado."
    )
    document.add_paragraph(
        "Si la mención puede representarse con un tramo continuo sin incluir palabras ajenas, use un span continuo. "
        "Si la expresión parece discontinua pero puede dividirse en menciones clínicas independientes, sepárelas solo "
        "si cada una tiene sentido propio. Si un span continuo incorporaría material que no pertenece a la mención y "
        "dividirla distorsionaría el significado, marque el caso para adjudicación o absténgase según el protocolo."
    )
    document.add_paragraph(
        "Esta decisión es gradual: los spans discontinuos pueden ser más fieles al texto, pero requieren una interfaz "
        "de selección múltiple, un formato de offsets diferente y reglas adicionales para solapamientos, adjudicación, "
        "métricas y modelos. La evidencia y las referencias completas están en Fundamentos metodológicos para anotadores "
        "SemantIAr: Alhassan et al. (2025), Dai et al. (2020) y Trivedi et al. (2020)."
    )
    document.add_paragraph(
        "Regla práctica: no incluya palabras intermedias que no pertenecen a la entidad solo para evitar una discontinuidad. "
        "Registrar el caso para revisión es preferible a alterar el literal o introducir ruido semántico."
    )
    document.add_heading("Sentido resuelto y cierre", level=2)
    document.add_paragraph(
        "Seleccione “Sentido resuelto” solo cuando haya elegido un significado disponible. Si no puede confirmar "
        "uno, use “Proponer sentido nuevo”, “Ambigua aun con contexto” o “No puedo determinarla”. Una decisión "
        "resuelta sin significado se normaliza a Pendiente al volver a cargar el archivo y bloquea el cierre de la nota."
    )
    document.add_paragraph(
        "Regla práctica: primero confirme el texto y el contexto; después registre la decisión léxica y el concepto clínico. "
        "No fuerce una resolución para completar una nota."
    )
    normalize_legacy_lexical_prompts(document)
    mark_docx_table_headers(document)
    add_word_navigation(document)
    document.save(DOCX_TARGET)


def normalize_legacy_lexical_prompts(document):
    """Keep the generated Word manual aligned with the current lexical UI."""
    for table in document.tables:
        for row in list(table.rows):
            values = [cell.text.strip() for cell in row.cells]
            if any(value == "¿Podés saber qué significa aquí?" for value in values):
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = run.text.replace(
                                "¿Podés saber qué significa aquí?", "¿Cómo se decide el significado?"
                            )
            if values and values[0] == "Comentario" and any(
                "¿Quedó algo importante sin registrar?" in value for value in values
            ):
                table._tbl.remove(row._tr)


def mark_docx_table_headers(document):
    for table in document.tables:
        if not table.rows or len(table.rows[0].cells) <= 1:
            continue
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))


def add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text, anchor):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend([properties, text_element])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_word_navigation(document):
    if "Related links" not in document.styles:
        style = document.styles.add_style("Related links", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        style.font.size = document.styles["Normal"].font.size

    headings = {paragraph.text: paragraph for paragraph in document.paragraphs if paragraph.text in WORD_BOOKMARKS}
    for bookmark_id, (heading, anchor) in enumerate(WORD_BOOKMARKS.items(), start=1000):
        paragraph = headings.get(heading)
        if paragraph is not None:
            add_bookmark(paragraph, anchor, bookmark_id)

    for heading, links in WORD_RELATED.items():
        heading_paragraph = headings.get(heading)
        if heading_paragraph is None:
            continue
        related = document.add_paragraph(style="Related links")
        related.add_run("También puede consultar: ").bold = True
        for index, (label, anchor) in enumerate(links):
            if index:
                related.add_run(" · ")
            add_internal_hyperlink(related, label, anchor)
        heading_paragraph._p.addnext(related._p)


if __name__ == "__main__":
    build_html()
    build_docx()
    print(f"OK: {HTML_TARGET.name}; {DOCX_TARGET.name}")
